import os
import subprocess
import mimetypes
import gc
import time
import threading
import queue
import json
import tempfile
import uuid
import numpy as np
import assemblyai as aai
import google.generativeai as genai
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAI
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Document loader imports
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredRTFLoader,
    DirectoryLoader,
    UnstructuredFileLoader
)

# Flask imports
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename

# Load environment variables
load_dotenv()

# API keys
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not ASSEMBLYAI_API_KEY or not GEMINI_API_KEY:
    raise ValueError("Missing API key(s) in environment variables.")

# Initialize AssemblyAI and Gemini
aai.settings.api_key = ASSEMBLYAI_API_KEY
genai.configure(api_key=GEMINI_API_KEY)

# ---------------------
# Performance Config
# ---------------------
MAX_WORKERS = 4  # Number of threads for parallel processing
BATCH_SIZE = 100  # Batch size for embeddings
EMBEDDER_CACHE_SIZE = 1024  # Size of embedding cache

# ---------------------
# Flask App Setup
# ---------------------
app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configure file upload settings
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx', 'doc', 'rtf', 'mp3', 'mp4'}
MAX_CONTENT_LENGTH = 100 * 1024 * 1024  # 100MB limit

# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Configure session storage
SESSION_FOLDER = os.path.join(os.getcwd(), 'sessions')
os.makedirs(SESSION_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Session storage for RAG instances
active_sessions = {}

# ---------------------
# Document Loaders
# ---------------------
def get_loader_mapping():
    """Return mapping of file extensions to document loaders."""
    return {
        ".pdf": PyPDFLoader,
        ".txt": TextLoader,
        ".docx": UnstructuredWordDocumentLoader,
        ".doc": UnstructuredWordDocumentLoader,
        ".rtf": UnstructuredRTFLoader,
        # Audio/video files are handled separately
        ".mp3": None,
        ".mp4": None,
        # Default loader for other file types
        "default": UnstructuredFileLoader
    }

def load_single_document(file_path: str) -> list:
    """Load a single document using appropriate loader"""
    file_ext = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)
    
    print(f"Processing file: {file_name} with extension {file_ext}")
    
    # Handle audio/video files with transcription
    if file_ext in ['.mp3', '.mp4']:
        try:
            transcription = transcribe_audio(file_path)
            doc = Document(page_content=transcription, metadata={'source_file': file_name})
            return [doc]
        except Exception as e:
            print(f"Error processing audio/video file {file_path}: {str(e)}")
            raise
    
    # Handle YouTube URLs (though this case should not normally reach here)
    elif "youtube.com" in file_path or "youtu.be" in file_path:
        try:
            audio_path = download_audio_from_youtube(file_path)
            transcription = transcribe_audio(audio_path)
            doc = Document(page_content=transcription, metadata={'source_file': os.path.basename(audio_path)})
            return [doc]
        except Exception as e:
            print(f"Error processing YouTube URL {file_path}: {str(e)}")
            raise
    
    # Handle regular files
    loader_mapping = get_loader_mapping()
    try:
        # Handle PDF files with a special case since it's most reliable
        if file_ext == '.pdf':
            print(f"Loading PDF file: {file_path}")
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif file_ext == '.txt':
            print(f"Loading TXT file: {file_path}")
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
            doc = Document(page_content=text_content, metadata={'source_file': file_name})
            documents = [doc]
        elif file_ext in ['.docx', '.doc']:
            print(f"Loading Word document: {file_path}")
            try:
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text])
                documents = [Document(page_content=text, metadata={'source_file': file_name})]
            except Exception as doc_e:
                print(f"Error with standard docx loader, trying UnstructuredWordDocumentLoader: {str(doc_e)}")
                # Fallback to UnstructuredWordDocumentLoader
                loader = UnstructuredWordDocumentLoader(file_path)
                documents = loader.load()
        elif file_ext in loader_mapping and loader_mapping[file_ext]:
            loader_class = loader_mapping[file_ext]
            print(f"Loading {file_ext} file with {loader_class.__name__}: {file_path}")
            loader = loader_class(file_path)
            documents = loader.load()
        else:
            # Try with default loader for unknown file types
            print(f"Attempting to load with default loader: {file_path}")
            loader = loader_mapping["default"](file_path)
            documents = loader.load()
        
        if not documents:
            print(f"Warning: No content extracted from {file_path}")
            # Create a minimal document with an error message
            documents = [Document(
                page_content=f"Error: No content could be extracted from this file: {file_name}",
                metadata={'source_file': file_name, 'error': 'No content extracted'}
            )]
        
        # Add source filename to metadata
        for doc in documents:
            if not hasattr(doc, 'metadata'):
                doc.metadata = {}
            doc.metadata['source_file'] = file_name
            # Ensure chunk identification
            if 'page' not in doc.metadata:
                doc.metadata['page'] = doc.metadata.get('page_number', 'N/A')
            
        print(f"Successfully loaded {len(documents)} document(s) from {file_path}")
        return documents
    except Exception as e:
        print(f"Error loading {file_path}: {str(e)}")
        # Instead of returning empty list, create a document with error information
        error_doc = Document(
            page_content=f"Error loading file: {file_name}\nError details: {str(e)}",
            metadata={'source_file': file_name, 'error': str(e)}
        )
        return [error_doc]

def load_multiple_documents(file_paths: list) -> list:
    """Load multiple documents from a list of file paths"""
    all_documents = []
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        # Submit load tasks
        future_to_path = {executor.submit(load_single_document, path): path for path in file_paths}
        
        # Process results as they complete
        for future in future_to_path:
            path = future_to_path[future]
            try:
                docs = future.result()
                all_documents.extend(docs)
            except Exception as e:
                print(f"Error processing {path}: {str(e)}")
    
    print(f"Loaded {len(all_documents)} total documents from {len(file_paths)} files")
    return all_documents

# ---------------------
# Embedding Helper with Caching
# ---------------------
@lru_cache(maxsize=1)
def get_embedder(lightweight=False, device="cpu"):
    """Create and cache embedders to prevent reloading models"""
    if lightweight:
        # Use a smaller, more efficient model for very large files
        return HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-MiniLM-L3-v2",
            model_kwargs={"device": device},
            encode_kwargs={"normalize_embeddings": False}
        )
    else:
        return HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-mpnet-base-v2",
            model_kwargs={"device": device},
            encode_kwargs={"normalize_embeddings": False}
        )

def make_embedder(lightweight=False):
    """Get cached embedder instance"""
    # Check if GPU is available
    device = "cuda" if os.environ.get("USE_GPU", "0") == "1" else "cpu"
    return get_embedder(lightweight, device)

# ---------------------
# YouTube + Audio Handling 
# ---------------------
def download_audio_from_youtube(url: str, output_path: str = None):
    try:
        # Create unique filename for the YouTube download
        if output_path is None:
            unique_id = str(uuid.uuid4())[:8]
            output_path = os.path.join(UPLOAD_FOLDER, f"youtube_{unique_id}.mp3")
        
        # Ensure yt-dlp is present (pip install yt-dlp)
        try:
            # Run yt-dlp with specific template to avoid filename issues
            cmd = [
                "yt-dlp", 
                "-x", 
                "--audio-format", "mp3", 
                "--audio-quality", "0",
                "-o", output_path,
                url
            ]
            print(f"Executing command: {' '.join(cmd)}")
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            print(f"✅ Downloaded audio to {output_path}")
            return output_path
        except subprocess.CalledProcessError as e:
            print(f"Error downloading YouTube video: {e}")
            print(f"Command output: {e.stdout}")
            print(f"Command error: {e.stderr}")
            raise ValueError(f"Failed to download YouTube video. Error: {str(e)}")
        except FileNotFoundError:
            raise ValueError("yt-dlp not installed. Please install with 'pip install yt-dlp'")
    except Exception as e:
        print(f"YouTube download failed: {str(e)}")
        raise

def transcribe_audio(file_path: str):
    try:
        transcript = aai.Transcriber().transcribe(file_path)
        print("✅ Transcription complete.")
        return transcript.text
    except Exception as e:
        print(f"Error during transcription: {str(e)}")
        raise ValueError(f"Failed to transcribe audio: {str(e)}")

# ---------------------
# Optimized Vector DB Creation
# ---------------------
def create_vector_db(documents: list, session_id: str, batch_size=BATCH_SIZE):
    """Create vector database from documents with optimized splitting and parallel embedding"""
    print(f"Processing {len(documents)} documents for vectorization")
    
    # Determine chunk size based on total document length
    total_text_length = sum(len(doc.page_content) for doc in documents)
    
    if total_text_length > 1000000:  # For very large documents
        chunk_size = 1500
        chunk_overlap = 150
    else:
        chunk_size = 500
        chunk_overlap = 50
    
    # Split text into chunks
    print(f"Splitting documents with chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    
    # Process each document separately to preserve metadata
    all_chunks = []
    for i, doc in enumerate(documents):
        chunks = splitter.split_documents([doc])
        # Add chunk number to each split
        for j, chunk in enumerate(chunks):
            chunk.metadata['chunk_number'] = j + 1  # 1-indexed for user-friendly reference
            chunk.metadata['total_chunks'] = len(chunks)
            # Make sure source file is preserved
            if 'source_file' not in chunk.metadata:
                chunk.metadata['source_file'] = doc.metadata.get('source_file', f'doc-{i}')
        all_chunks.extend(chunks)
    
    chunks = all_chunks
    print(f"✅ Split into {len(chunks)} chunks.")
    
    # Use lightweight embedder for very large document sets
    use_lightweight = len(chunks) > 500
    embedder = make_embedder(lightweight=use_lightweight)
    
    # Process in parallel batches for large document sets
    if len(chunks) > 200:
        db = create_vector_db_parallel(chunks, embedder, batch_size)
    else:
        db = FAISS.from_documents(chunks, embedder)
    
    # Save the database to the session folder
    os.makedirs(os.path.join(SESSION_FOLDER, session_id), exist_ok=True)
    db.save_local(os.path.join(SESSION_FOLDER, session_id, "faiss_index"))
    print(f"✅ FAISS vector DB created for session {session_id}.")
    return db

def create_vector_db_parallel(documents, embedder, batch_size=BATCH_SIZE):
    """Process embeddings in parallel batches for better performance"""
    print(f"🔄 Processing embeddings in parallel batches of {batch_size}")
    
    # Split into manageable batches
    total_batches = (len(documents) - 1) // batch_size + 1
    batches = [documents[i:min(i+batch_size, len(documents))] 
               for i in range(0, len(documents), batch_size)]
    
    # Create queue for results
    db_queue = queue.Queue()
    
    # Process first batch separately to initialize the DB
    print(f"  ⏳ Processing batch 1/{total_batches}")
    db = FAISS.from_documents(batches[0], embedder)
    
    if total_batches > 1:
        # Process remaining batches with a thread pool
        def process_batch(batch_idx):
            batch = batches[batch_idx]
            print(f"  ⏳ Processing batch {batch_idx+1}/{total_batches} ({len(batch)} chunks)")
            try:
                temp_db = FAISS.from_documents(batch, embedder)
                db_queue.put((batch_idx, temp_db))
            except Exception as e:
                print(f"Error processing batch {batch_idx+1}: {e}")
                db_queue.put((batch_idx, None))
        
        # Use a thread pool with limited workers to avoid memory exhaustion
        with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, total_batches-1)) as executor:
            # Submit all batches after the first one
            futures = [executor.submit(process_batch, i) for i in range(1, total_batches)]
        
        # Process results in order
        while not db_queue.empty():
            _, temp_db = db_queue.get()
            if temp_db:
                db.merge_from(temp_db)
            # Force cleanup after each merge
            gc.collect()
    
    return db

def load_vector_db(session_id: str):
    """Load existing vector database from disk for a session"""
    session_path = os.path.join(SESSION_FOLDER, session_id, "faiss_index")
    if os.path.exists(session_path):
        embedder = make_embedder()
        return FAISS.load_local(session_path, embedder, allow_dangerous_deserialization=True)
    return None

# ---------------------
# QA Chain (RAG)
# ---------------------
def create_qa_chain(db):
    """Create RAG query chain"""
    retriever = db.as_retriever(
        search_type="mmr",  # Use Maximum Marginal Relevance for better diversity
        search_kwargs={"k": 5, "fetch_k": 15, "lambda_mult": 0.7},
    )
    system_prompt = """
Based on the following information from the documents:{context}

Please answer this question: {input}

IMPORTANT INSTRUCTIONS FOR CITATIONS:
1. ONLY cite information that comes directly from the provided documents.
2. When referencing information FROM THE DOCUMENTS, use numbered citations like [1], [2], etc.
3. DO NOT use citations when providing general knowledge that isn't in the documents.
4. After your answer, include a "REFERENCES:" section that lists all sources you cited.
5. Each reference should correspond to a specific document chunk and follow this format:
   [number] -> Document: [filename], Chunk: [chunk_number]
   If page information is available, include it: [number] -> Document: [filename], Page: [page], Chunk: [chunk_number]

If the provided documents don't contain information related to the question:
1. Clearly state "The provided documents don't contain information about this topic."
2. You may then provide a brief general answer WITHOUT using any citations.
3. Do not fabricate citations for general knowledge.
"""

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        ("human", "{input}")
    ])

    llm = GoogleGenerativeAI(
        model="gemini-2.0-flash",
        google_api_key=GEMINI_API_KEY,
        temperature=0.3
    )

    doc_chain = create_stuff_documents_chain(llm=llm, prompt=prompt)
    return create_retrieval_chain(retriever, doc_chain)

# ---------------------
# Optimized Gemini Analysis
# ---------------------
def analyze_with_gemini(text: str, question: str = "What is this about?"):
    """Analyze text with Gemini"""
    # For very large texts, use a more efficient summarization approach
    if len(text) > 100000:
        print("⚠️ Text too large, summarizing first part and sampling throughout")
        # Take first 25K, middle 25K and last 25K for a more balanced view
        first = text[:25000]
        middle_start = max(0, len(text)//2 - 12500)
        middle = text[middle_start:middle_start+25000]
        last = text[-25000:] if len(text) > 25000 else ""
        analysis_text = f"{first}\n\n[...middle content omitted...]\n\n{middle}\n\n[...additional content omitted...]\n\n{last}"
    else:
        analysis_text = text
        
    model = genai.GenerativeModel("gemini-2.0-flash", generation_config={"temperature": 0.2})
    prompt = f"""Here is the content, include the citation of the source in the answer [source filename]:

{analysis_text}

Now, answer:
{question}
"""
    response = model.generate_content(prompt)
    print("✅ Gemini summary generated.")
    return response.text.strip()

# ---------------------
# Process Documents with Session Management
# ---------------------
def process_documents(file_paths: list, session_id: str, initial_question: str = "What is this about?"):
    """Process one or more documents and return session details"""
    start_time = time.time()
    
    # Load documents
    print(f"📖 Loading content from {len(file_paths)} file(s)...")
    documents = []
    error_messages = []
    
    # Try to load each document, collect any errors
    for path in file_paths:
        try:
            docs = load_single_document(path)
            if docs:
                documents.extend(docs)
            else:
                error_messages.append(f"No content extracted from {os.path.basename(path)}")
        except Exception as e:
            error_messages.append(f"Error loading {os.path.basename(path)}: {str(e)}")
    
    if not documents:
        error_detail = "; ".join(error_messages)
        raise ValueError(f"Failed to extract content from any of the provided files. Details: {error_detail}")
    
    # Create combined content preview
    combined_content = "\n".join([doc.page_content[:500] for doc in documents[:3]])
    if len(documents) > 3:
        combined_content += f"\n\n... and {len(documents) - 3} more document(s) ..."
    combined_content = combined_content[:1000] + ("..." if len(combined_content) > 1000 else "")
    
    # Calculate total content length
    total_content_length = sum(len(doc.page_content) for doc in documents)
    
    # Process content
    print(f"🔍 Creating vector database for session {session_id}...")
    db = create_vector_db(documents, session_id)
    
    # Create QA chain
    qa_chain = create_qa_chain(db)
    
    # Generate initial summary
    # For summary, use content from first document if multiple files
    sample_text = documents[0].page_content
    if total_content_length > 100000:
        # If content is very large, limit sample text
        sample_text = sample_text[:100000]
    summary = analyze_with_gemini(sample_text, initial_question)
    
    # Answer initial question
    initial_answer = qa_chain.invoke({"input": initial_question}).get("answer", "No answer found")
    
    # Calculate processing time
    processing_time = time.time() - start_time
    
    # Get list of file names
    file_names = list(set(doc.metadata.get('source_file', 'unknown') for doc in documents))
    
    # Store QA chain in active sessions
    active_sessions[session_id] = {
        "qa_chain": qa_chain,
        "content_preview": combined_content,
        "content_length": total_content_length,
        "file_paths": file_paths,
        "file_names": file_names,
        "created_at": time.time(),
        "last_accessed": time.time()
    }
    
    # Return results
    return {
        "session_id": session_id,
        "file_names": file_names,
        "num_files": len(file_paths),
        "content_length": total_content_length,
        "content_preview": combined_content,
        "summary": summary,
        "initial_answer": initial_answer,
        "processing_time": round(processing_time, 2)
    }

# ---------------------
# Flask Routes
# ---------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/health', methods=['GET'])
def health_check():
    """API health check endpoint"""
    return jsonify({
        "status": "ok",
        "timestamp": time.time(),
        "version": "1.0.0"
    })

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Upload one or more files for processing"""
    print('Uploading files...')
    
    # Check if file(s) or URL are present in request
    has_files = 'file' in request.files
    has_multiple_files = request.files.getlist('file') and len(request.files.getlist('file')) > 1
    print(f"Files found: {len(request.files.getlist('file'))}")
    
    # Check for YouTube URL in form data or JSON payload
    youtube_url = ''
    if request.form and 'youtube_url' in request.form:
        youtube_url = request.form['youtube_url']
        print(f"Found YouTube URL in form data: {youtube_url}")
    elif request.is_json and request.json and 'youtube_url' in request.json:
        youtube_url = request.json['youtube_url']
        print(f"Found YouTube URL in JSON payload: {youtube_url}")
    elif request.data:
        try:
            json_data = json.loads(request.data)
            if 'youtube_url' in json_data:
                youtube_url = json_data['youtube_url']
                print(f"Found YouTube URL in raw request data: {youtube_url}")
        except:
            pass
            
    has_youtube_url = bool(youtube_url)
    print(f"YouTube URL found: {has_youtube_url}")
    
    if not has_files and not has_youtube_url:
        print("Neither files nor YouTube URL found in request")
        print(f"Form data keys: {list(request.form.keys()) if request.form else 'None'}")
        print(f"JSON data: {request.json if request.is_json else 'None'}")
        return jsonify({"error": "No file or YouTube URL provided"}), 400
    
    # Generate a session ID
    session_id = str(uuid.uuid4())
    file_paths = []
    
    try:
        # Process file uploads
        if has_files:
            files = request.files.getlist('file')  # Get multiple files
            
            for file in files:
                # Check if valid file
                if file.filename == '':
                    continue
                    
                if not allowed_file(file.filename):
                    return jsonify({"error": f"File type not allowed for {file.filename}. Supported types: {', '.join(ALLOWED_EXTENSIONS)}"}), 400
                
                # Save the file
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{session_id}_{filename}")
                file.save(file_path)
                file_paths.append(file_path)
                print(f"Saved file: {file_path}")
                
        # Process YouTube URL
        if has_youtube_url:
            print(f"Processing YouTube URL: {youtube_url}")
            
            # Handle YouTube URLs
            if "youtube.com" in youtube_url or "youtu.be" in youtube_url:
                try:
                    file_path = download_audio_from_youtube(youtube_url)
                    file_paths.append(file_path)
                    print(f"YouTube processing complete: {file_path}")
                except Exception as e:
                    return jsonify({"error": f"YouTube processing failed: {str(e)}"}), 500
            else:
                return jsonify({"error": "Only YouTube URLs are supported"}), 400
        
        if not file_paths:
            return jsonify({"error": "No valid files provided"}), 400
        
        # Get initial question if provided - check multiple sources
        initial_question = None
        
        # Try to get from form data first
        if request.form and 'question' in request.form:
            initial_question = request.form.get('question')
        # Then try from JSON if available
        elif request.is_json and request.json and 'question' in request.json:
            initial_question = request.json.get('question')
        # Finally try from raw data as JSON
        elif request.data:
            try:
                json_data = json.loads(request.data)
                if 'question' in json_data:
                    initial_question = json_data['question']
            except:
                pass
        
        # Default if not found
        if not initial_question:
            initial_question = "What is this about?"
            
        print(f"Using initial question: {initial_question}")
        
        # Process documents in a background thread to avoid blocking
        def process_in_background():
            try:
                result = process_documents(file_paths, session_id, initial_question)
                # Store result for later retrieval
                active_sessions[session_id]["processing_result"] = result
                active_sessions[session_id]["processing_complete"] = True
            except Exception as e:
                print(f"Error processing documents: {str(e)}")
                active_sessions[session_id]["processing_error"] = str(e)
                active_sessions[session_id]["processing_complete"] = True
        
        # Initialize session entry
        active_sessions[session_id] = {
            "file_paths": file_paths,
            "created_at": time.time(),
            "processing_complete": False
        }
        
        # Start processing in background
        threading.Thread(target=process_in_background).start()
        
        # Return session ID immediately
        return jsonify({
            "session_id": session_id,
            "status": "processing",
            "message": f"{'Files' if has_multiple_files else 'File'} uploaded and processing started"
        })
        
    except Exception as e:
        print(f"Upload error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions/<session_id>/status', methods=['GET'])
def session_status(session_id):
    """Check the status of a processing session"""
    if session_id not in active_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    session = active_sessions[session_id]
    
    # Update last accessed time
    session["last_accessed"] = time.time()
    
    # Return status
    status_info = {
        "session_id": session_id,
        "processing_complete": session.get("processing_complete", False),
        "created_at": session.get("created_at")
    }
    
    # Add error if present
    if "processing_error" in session:
        status_info["error"] = session["processing_error"]
    
    # Add result summary if processing is complete
    if session.get("processing_complete", False) and "processing_error" not in session:
        result = session.get("processing_result", {})
        status_info.update({
            "file_names": session.get("file_names", []),
            "num_files": len(session.get("file_paths", [])),
            "content_length": result.get("content_length", 0),
            "processing_time": result.get("processing_time", 0)
        })
    
    return jsonify(status_info)

@app.route('/api/sessions/<session_id>/result', methods=['GET'])
def session_result(session_id):
    """Get the complete processing result for a session"""
    if session_id not in active_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    session = active_sessions[session_id]
    
    # Update last accessed time
    session["last_accessed"] = time.time()
    
    # Check if processing is complete
    if not session.get("processing_complete", False):
        return jsonify({
            "session_id": session_id,
            "status": "processing",
            "message": "Processing not yet complete"
        }), 202
    
    # Check for errors
    if "processing_error" in session:
        return jsonify({
            "session_id": session_id,
            "status": "error",
            "error": session["processing_error"]
        }), 500
    
    # Return full result
    return jsonify(session.get("processing_result", {}))

@app.route('/api/sessions/<session_id>/query', methods=['POST'])
def query_session(session_id):
    """Query a processed document"""
    if session_id not in active_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    session = active_sessions[session_id]
    
    # Update last accessed time
    session["last_accessed"] = time.time()
    
    # Check if processing is complete
    if not session.get("processing_complete", False):
        return jsonify({
            "status": "processing",
            "message": "Document still processing"
        }), 202
    
    # Check for errors
    if "processing_error" in session:
        return jsonify({
            "status": "error",
            "error": session["processing_error"]
        }), 500
    
    # Get query from request
    data = request.json
    if not data or "query" not in data:
        return jsonify({"error": "No query provided"}), 400
    
    query = data["query"]
    detailed = data.get("detailed", False)
    
    try:
        # Get QA chain
        qa_chain = session.get("qa_chain")
        if not qa_chain:
            return jsonify({"error": "Session QA chain not found"}), 500
        
        # Execute query
        start_time = time.time()
        result = qa_chain.invoke({"input": query})
        processing_time = time.time() - start_time
        
        # Prepare response
        answer = result.get("answer", "No answer found")
        response = {
            "query": query,
            "answer": answer,
            "processing_time": round(processing_time, 2)
        }
        
        # Add source documents if detailed requested
        if detailed and "source_documents" in result:
            sources = []
            for i, doc in enumerate(result["source_documents"][:5]):  # Include top 5 sources
                sources.append({
                    "id": i + 1,
                    "content": doc.page_content,
                    "metadata": doc.metadata
                })
            response["sources"] = sources
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions/<session_id>', methods=['DELETE'])
def delete_session(session_id):
    """Delete a session and its associated files"""
    if session_id not in active_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    try:
        # Get file path from session
        file_paths = active_sessions[session_id].get("file_paths", [])
        
        # Delete the files if they exist
        for file_path in file_paths:
            if os.path.exists(file_path):
                os.remove(file_path)
        
        # Delete vector database directory
        session_dir = os.path.join(SESSION_FOLDER, session_id)
        if os.path.exists(session_dir):
            import shutil
            shutil.rmtree(session_dir)
            
        # Remove session from active sessions
        del active_sessions[session_id]
        
        # Trigger garbage collection
        gc.collect()
        
        return jsonify({"status": "success", "message": "Session deleted"})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions', methods=['GET'])
def list_sessions():
    """List all active sessions"""
    sessions = []
    for session_id, session_data in active_sessions.items():
        sessions.append({
            "session_id": session_id,
            "file_names": session_data.get("file_names", []),
            "num_files": len(session_data.get("file_paths", [])),
            "created_at": session_data.get("created_at"),
            "last_accessed": session_data.get("last_accessed"),
            "processing_complete": session_data.get("processing_complete", False),
            "has_error": "processing_error" in session_data
        })
    
    return jsonify({"sessions": sessions})

# Periodic cleanup of old sessions
def cleanup_old_sessions():
    """Remove sessions that haven't been accessed for a while"""
    current_time = time.time()
    session_timeout = 3600  # 1 hour
    
    sessions_to_remove = []
    for session_id, session_data in active_sessions.items():
        if current_time - session_data.get("last_accessed", current_time) > session_timeout:
            sessions_to_remove.append(session_id)
    
    for session_id in sessions_to_remove:
        try:
            # Get file path from session
            file_paths = active_sessions[session_id].get("file_paths", [])
            
            # Delete the files if they exist
            for file_path in file_paths:
                if os.path.exists(file_path):
                    os.remove(file_path)
            
            # Delete vector database directory
            session_dir = os.path.join(SESSION_FOLDER, session_id)
            if os.path.exists(session_dir):
                import shutil
                shutil.rmtree(session_dir)
                
            # Remove session from active sessions
            del active_sessions[session_id]
            
            print(f"Cleaned up session {session_id} due to inactivity")
            
        except Exception as e:
            print(f"Error cleaning up session {session_id}: {str(e)}")
    
    # Schedule next cleanup
    threading.Timer(600, cleanup_old_sessions).start()  # Run every 10 minutes

# ---------------------
# Main entry point
# ---------------------
if __name__ == "__main__":
    print("🚀 Starting RAG API Server")
    # Start cleanup thread
    cleanup_old_sessions()
    # Start the server
    app.run(host='0.0.0.0', port=5000, debug=False)